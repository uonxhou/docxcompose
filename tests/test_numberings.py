from docx import Document
from docxcompose.composer import Composer
from docxcompose.utils import xpath
from utils import ComparableDocument
from utils import ComposedComparableDocument
from utils import docx_path
import pytest


def test_abstractnums_from_styles_are_not_duplicated(multiple_numberings):
    anums = xpath(
        multiple_numberings.doc.part.numbering_part.element,
        './/w:abstractNum[.//w:pStyle]')
    assert len(anums) == 2


def test_restart_first_numbering(multiple_numberings):
    paragraphs = xpath(multiple_numberings.doc.element.body, './/w:p')
    assert len(xpath(paragraphs[9], './/w:numId')) == 1


def test_do_not_restart_numbering_of_bullets(mixed_numberings):
    paragraphs = xpath(mixed_numberings.doc.element.body, './/w:p')
    assert len(xpath(paragraphs[10], './/w:numId')) == 0


def test_preserve_zero_numbering_references(numberings_with_zero_reference):
    numberings_with_zero_ref = xpath(
        numberings_with_zero_reference.doc.element.body,
        './/w:p//w:numId[@w:val="0"]')
    assert len(numberings_with_zero_ref) == 2


def test_restart_numberings():
    doc = ComparableDocument(
        Document(docx_path("composed/numberings_restart.docx")))
    composed = ComposedComparableDocument(
        "numberings_restart.docx", "numberings_restart.docx")
    assert composed == doc


@pytest.fixture
def numberings_with_zero_reference():
    composer = Composer(Document(
        docx_path("numbering_reference_to_numbering_zero.docx")))
    composer.append(Document(
        docx_path("numbering_reference_to_numbering_zero.docx")))
    return composer


@pytest.fixture
def numberings_in_styles():
    composer = Composer(Document(docx_path("master.docx")))
    composer.append(Document(docx_path("numberings_styles.docx")))
    return composer


@pytest.fixture
def multiple_numberings():
    composer = Composer(Document(docx_path("numberings_styles.docx")))
    composer.append(Document(docx_path("numberings_styles.docx")))
    return composer


@pytest.fixture
def mixed_numberings():
    composer = Composer(Document(docx_path("numberings_restart.docx")))
    composer.append(Document(docx_path("numberings_restart.docx")))
    return composer
